import sys
import argparse
from docx import Document


def load_docx(file_path):
    """Load a DOCX file and return the Document object."""
    return Document(file_path)


def save_docx(document, file_path):
    """Save the Document object to a DOCX file."""
    document.save(file_path)


def replace_text(document, old_text, new_text):
    """Replace all occurrences of old_text with new_text in the document."""
    for para in document.paragraphs:
        if old_text in para.text:
            inline = para.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    inline[i].text = inline[i].text.replace(old_text, new_text)
    # Also handle tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text(cell, old_text, new_text)


def add_paragraph(document, text, style=None, indent=None, spacing=None, font_props=None):
    """Add a new paragraph with advanced formatting options."""
    para = document.add_paragraph(text)
    if style:
        para.style = style
    if indent:
        if 'left' in indent:
            para.paragraph_format.left_indent = indent['left']
        if 'right' in indent:
            para.paragraph_format.right_indent = indent['right']
        if 'first_line' in indent:
            para.paragraph_format.first_line_indent = indent['first_line']
    if spacing:
        if 'before' in spacing:
            para.paragraph_format.space_before = spacing['before']
        if 'after' in spacing:
            para.paragraph_format.space_after = spacing['after']
        if 'line' in spacing:
            para.paragraph_format.line_spacing = spacing['line']
    if font_props:
        run = para.runs[0] if para.runs else para.add_run()
        if 'name' in font_props:
            run.font.name = font_props['name']
        if 'size' in font_props:
            from docx.shared import Pt
            run.font.size = Pt(font_props['size'])
        if 'bold' in font_props:
            run.font.bold = font_props['bold']
        if 'italic' in font_props:
            run.font.italic = font_props['italic']
        if 'underline' in font_props:
            run.font.underline = font_props['underline']
        if 'color' in font_props:
            from docx.shared import RGBColor
            run.font.color.rgb = RGBColor(*font_props['color'])
    return para


def remove_paragraphs_with_keyword(document, keyword):
    """Remove all paragraphs containing the specified keyword."""
    paras_to_remove = [para for para in document.paragraphs if keyword in para.text]
    for para in paras_to_remove:
        p = para._element
        p.getparent().remove(p)
        p._p = p._element = None


def parse_indent(indent_str):
    # Format: left,right,first_line (in points)
    if not indent_str:
        return None
    vals = indent_str.split(',')
    keys = ['left', 'right', 'first_line']
    from docx.shared import Pt
    return {k: Pt(float(v)) for k, v in zip(keys, vals) if v}

def parse_spacing(spacing_str):
    # Format: before,after,line (in points, except line can be float)
    if not spacing_str:
        return None
    vals = spacing_str.split(',')
    keys = ['before', 'after', 'line']
    from docx.shared import Pt
    result = {}
    if vals[0]:
        result['before'] = Pt(float(vals[0]))
    if len(vals) > 1 and vals[1]:
        result['after'] = Pt(float(vals[1]))
    if len(vals) > 2 and vals[2]:
        try:
            result['line'] = float(vals[2])
        except ValueError:
            pass
    return result

def parse_font(font_str):
    # Format: name,size,bold,italic,underline,colorRRGGBB
    if not font_str:
        return None
    vals = font_str.split(',')
    props = {}
    if vals[0]:
        props['name'] = vals[0]
    if len(vals) > 1 and vals[1]:
        props['size'] = float(vals[1])
    if len(vals) > 2 and vals[2]:
        props['bold'] = vals[2].lower() == 'true'
    if len(vals) > 3 and vals[3]:
        props['italic'] = vals[3].lower() == 'true'
    if len(vals) > 4 and vals[4]:
        props['underline'] = vals[4].lower() == 'true'
    if len(vals) > 5 and vals[5]:
        hex_color = vals[5]
        props['color'] = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    return props

def add_toc_placeholder(document):
    """Add a Table of Contents placeholder (Word will generate ToC when opened)."""
    para = document.add_paragraph()
    run = para.add_run('Table of Contents')
    run.bold = True
    para.alignment = 1  # Center
    document.add_paragraph('')  # Blank line
    # Add field code for ToC (Word will update)
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p = document.add_paragraph()
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'TOC \o "1-3" \h \z \u')
    p._p.append(fldSimple)


def add_heading_with_number(document, text, level=1):
    """Insert a heading with automatic numbering."""
    heading_text = f"{level}. {text}" if level else text
    para = document.add_paragraph(heading_text, style=f'Heading {level}')
    return para


def add_caption(document, caption_text, label='Figure'):
    """Insert a caption for a figure or table."""
    para = document.add_paragraph(f'{label}: {caption_text}')
    para.style = 'Caption'
    return para


def add_page_break(document):
    """Insert a page break."""
    document.add_page_break()


def set_page_layout(document, orientation='portrait', margins=(2.54,2.54,2.54,2.54)):
    """Set page orientation and margins (in cm)."""
    from docx.shared import Cm
    section = document.sections[0]
    if orientation == 'landscape':
        from docx.enum.section import WD_ORIENT
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(margins[0])
    section.right_margin = Cm(margins[1])
    section.bottom_margin = Cm(margins[2])
    section.left_margin = Cm(margins[3])


def add_header_footer_with_page_number(document, header_text=None, footer_text=None):
    section = document.sections[0]
    if header_text:
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = header_text
    if footer_text:
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = footer_text
    # Add page number field to footer
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    footer = section.footer
    para = footer.add_paragraph()
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE')
    para._p.append(fldSimple)
    para.alignment = 1  # Center


def main():
    parser = argparse.ArgumentParser(description='DOCX Modifier Script with Advanced Formatting and Academic Features')
    parser.add_argument('input', help='Input DOCX file path')
    parser.add_argument('output', help='Output DOCX file path')
    parser.add_argument('--replace', nargs=2, metavar=('OLD', 'NEW'), help='Replace all OLD text with NEW text')
    parser.add_argument('--add-paragraph', metavar='TEXT', help='Add a paragraph with TEXT')
    parser.add_argument('--para-style', metavar='STYLE', help='Style for new paragraph (e.g., Normal, Heading1)')
    parser.add_argument('--indent', metavar='LEFT,RIGHT,FIRST', help='Indentation in points, e.g., 36,0,18')
    parser.add_argument('--spacing', metavar='BEFORE,AFTER,LINE', help='Spacing in points, e.g., 12,12,1.5')
    parser.add_argument('--font', metavar='NAME,SIZE,BOLD,ITALIC,UNDERLINE,COLOR', help='Font properties, e.g., Arial,12,True,False,False,FF0000')
    parser.add_argument('--remove-paragraphs', metavar='KEYWORD', help='Remove paragraphs containing KEYWORD')
    # Academic features
    parser.add_argument('--add-toc', action='store_true', help='Insert Table of Contents placeholder')
    parser.add_argument('--add-heading', nargs=2, metavar=('LEVEL', 'TEXT'), help='Add heading with numbering (LEVEL: 1-3)')
    parser.add_argument('--add-caption', nargs=2, metavar=('LABEL', 'TEXT'), help='Add caption (LABEL: Figure/Table)')
    parser.add_argument('--add-page-break', action='store_true', help='Insert a page break')
    parser.add_argument('--set-page-layout', metavar='ORIENTATION,MARGINS', help='Set page layout: orientation (portrait/landscape), margins (cm, comma-separated, e.g., 2.54,2.54,2.54,2.54)')
    parser.add_argument('--header', metavar='TEXT', help='Add header text')
    parser.add_argument('--footer', metavar='TEXT', help='Add footer text')
    args = parser.parse_args()

    doc = load_docx(args.input)

    if args.replace:
        old, new = args.replace
        replace_text(doc, old, new)
        print(f"Replaced all '{old}' with '{new}'.")

    if args.add_paragraph:
        indent = parse_indent(args.indent)
        spacing = parse_spacing(args.spacing)
        font_props = parse_font(args.font)
        add_paragraph(doc, args.add_paragraph, style=args.para_style, indent=indent, spacing=spacing, font_props=font_props)
        print(f"Added new paragraph: {args.add_paragraph}")

    if args.remove_paragraphs:
        remove_paragraphs_with_keyword(doc, args.remove_paragraphs)
        print(f"Removed paragraphs containing: {args.remove_paragraphs}")

    if args.add_toc:
        add_toc_placeholder(doc)
        print("Added Table of Contents placeholder.")

    if args.add_heading:
        level, text = args.add_heading
        add_heading_with_number(doc, text, int(level))
        print(f"Added heading level {level}: {text}")

    if args.add_caption:
        label, caption_text = args.add_caption
        add_caption(doc, caption_text, label=label)
        print(f"Added caption: {label}: {caption_text}")

    if args.add_page_break:
        add_page_break(doc)
        print("Added page break.")

    if args.set_page_layout:
        orientation, margins_str = args.set_page_layout.split(',')[0], ','.join(args.set_page_layout.split(',')[1:])
        margins = tuple(float(x) for x in margins_str.split(',')) if margins_str else (2.54,2.54,2.54,2.54)
        set_page_layout(doc, orientation=orientation, margins=margins)
        print(f"Set page layout: orientation={orientation}, margins={margins}")

    if args.header or args.footer:
        add_header_footer_with_page_number(doc, header_text=args.header, footer_text=args.footer)
        print("Added header/footer with page number.")

    save_docx(doc, args.output)
    print(f"Saved modified document to {args.output}")


if __name__ == '__main__':
    main()
