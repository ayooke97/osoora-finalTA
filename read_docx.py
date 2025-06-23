import sys
import time
from docx import Document

from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_paragraph_advanced_styles(para):
    styles = {}
    # Alignment
    align_map = {
        WD_ALIGN_PARAGRAPH.LEFT: 'Left',
        WD_ALIGN_PARAGRAPH.CENTER: 'Center',
        WD_ALIGN_PARAGRAPH.RIGHT: 'Right',
        WD_ALIGN_PARAGRAPH.JUSTIFY: 'Justify',
    }
    styles['Alignment'] = align_map.get(para.alignment, 'Default')
    # Indentation
    styles['Indent Left'] = para.paragraph_format.left_indent.pt if para.paragraph_format.left_indent else 'Default'
    styles['Indent Right'] = para.paragraph_format.right_indent.pt if para.paragraph_format.right_indent else 'Default'
    styles['Indent First Line'] = para.paragraph_format.first_line_indent.pt if para.paragraph_format.first_line_indent else 'Default'
    # Spacing
    styles['Space Before'] = para.paragraph_format.space_before.pt if para.paragraph_format.space_before else 'Default'
    styles['Space After'] = para.paragraph_format.space_after.pt if para.paragraph_format.space_after else 'Default'
    styles['Line Spacing'] = para.paragraph_format.line_spacing if para.paragraph_format.line_spacing else 'Default'
    # Outline level (heading level)
    styles['Outline Level'] = para.style.base_style.name if para.style and para.style.base_style else 'Default'
    # Borders and shading (not directly supported by python-docx)
    styles['Borders'] = 'Not Supported by python-docx'
    styles['Shading'] = 'Not Supported by python-docx'
    return styles

def get_run_advanced_styles(run):
    font = run.font
    adv = {}
    adv['All Caps'] = 'Yes' if font.all_caps else 'No'
    adv['Small Caps'] = 'Yes' if font.small_caps else 'No'
    adv['Strike'] = 'Yes' if font.strike else 'No'
    adv['Double Strike'] = 'Yes' if font.double_strike else 'No'
    adv['Highlight Color'] = font.highlight_color if font.highlight_color else 'None'
    adv['Char Style'] = run.style.name if run.style else 'Default'
    return adv

def read_docx(file_path):
    doc = Document(file_path)
    for para_idx, para in enumerate(doc.paragraphs):
        print(f"\nParagraph {para_idx+1} (Style: {para.style.name if para.style else 'None'}):")
        adv_styles = get_paragraph_advanced_styles(para)
        for k, v in adv_styles.items():
            print(f"    {k}: {v}")
        for run_idx, run in enumerate(para.runs):
            font = run.font
            # Font properties
            font_name = font.name if font.name else 'Default'
            font_size = font.size.pt if font.size else 'Default'
            bold = 'Yes' if run.bold else 'No'
            italic = 'Yes' if run.italic else 'No'
            underline = 'Yes' if run.underline else 'No'
            color = font.color.rgb if font.color and font.color.rgb else 'Default'
            print(f"  Run {run_idx+1}:")
            print(f"    Text      : {run.text}")
            print(f"    Font Name : {font_name}")
            print(f"    Font Size : {font_size}")
            print(f"    Bold      : {bold}")
            print(f"    Italic    : {italic}")
            print(f"    Underline : {underline}")
            print(f"    Color     : {color}")
            # Advanced run styles
            adv = get_run_advanced_styles(run)
            for k, v in adv.items():
                print(f"    {k}: {v}")

import zipfile
import os

def extract_images_from_docx(file_path, output_dir=None):
    """
    Extract all images from a DOCX file and save them to the specified output directory.
    If output_dir is None, saves to the current directory.
    """
    if output_dir is None:
        output_dir = os.getcwd()
    with zipfile.ZipFile(file_path, 'r') as docx_zip:
        image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
        if not image_files:
            print("No images found in the DOCX file.")
            return
        for image_name in image_files:
            image_data = docx_zip.read(image_name)
            image_filename = os.path.basename(image_name)
            output_path = os.path.join(output_dir, image_filename)
            with open(output_path, 'wb') as img_file:
                img_file.write(image_data)
            print(f"Extracted: {output_path}")

def extract_content_and_style(file_path, chunk_size=20):
    # For clean output - prevents text overlap
    import io
    import sys
    original_stdout = sys.stdout
    
    doc = Document(file_path)
    total_paras = len(doc.paragraphs)
    output_lines = []
    
    # Always clean the output file first
    with open('output_chunk_bakti.txt', 'w', encoding='utf-8') as f:
        f.write("")
    
    for chunk_start in range(0, total_paras, chunk_size):
        chunk_end = min(chunk_start + chunk_size, total_paras)
        chunk_lines = []
        for para_idx in range(chunk_start, chunk_end):
            para = doc.paragraphs[para_idx]
            # Consistent format as in output_chunk.txt
            chunk_lines.append(f"Paragraph {para_idx+1}")
            chunk_lines.append(f"  Text: {para.text}")
            chunk_lines.append(f"  Style: {para.style.name if para.style else 'None'}")
            adv_styles = get_paragraph_advanced_styles(para)
            for k, v in adv_styles.items():
                chunk_lines.append(f"    {k}: {v}")
            # Process runs
            for run_idx, run in enumerate(para.runs):
                font = run.font
                # Ensure consistent, clean output
                font_name = font.name if font.name else 'Default'
                font_size = font.size.pt if font.size else 'Default'
                bold = 'Yes' if run.bold else 'No'
                italic = 'Yes' if run.italic else 'No'
                underline = 'Yes' if run.underline else 'No'
                color = font.color.rgb if font.color and font.color.rgb else 'Default'
                chunk_lines.append(f"    Run {run_idx+1}:")
                chunk_lines.append(f"      Text      : {run.text}")
                chunk_lines.append(f"      Font Name : {font_name}")
                chunk_lines.append(f"      Font Size : {font_size}")
                chunk_lines.append(f"      Bold      : {bold}")
                chunk_lines.append(f"      Italic    : {italic}")
                chunk_lines.append(f"      Underline : {underline}")
                chunk_lines.append(f"      Color     : {color}")
                adv = get_run_advanced_styles(run)
                for k, v in adv.items():
                    chunk_lines.append(f"      {k}: {v}")
        
        # Collect output in master list
        output_lines.extend(chunk_lines)
        
        # Use string buffer for clean output with no overlap
        buffer = io.StringIO()
        buffer.write("\n".join(chunk_lines))
        clean_output = buffer.getvalue()
        buffer.close()
        
        # Print current chunk cleanly
        print(clean_output)
        
        # Write to file with clean formatting
        with open('output_chunk_bakti.txt', 'a', encoding='utf-8') as f:
            f.write(clean_output + "\n\n")
        
        # Handle pausing between chunks
        if chunk_end < total_paras:
            # Get auto_stream from args or default to True
            auto_stream = getattr(args, 'auto_stream', True) if 'args' in globals() else True
            
            if auto_stream:
                # Auto-streaming mode: brief delay between chunks
                time.sleep(0.5)
            else:
                try:
                    # Try interactive pause
                    if sys.stdin.isatty():
                        input(f"\n--- Showing paragraphs {chunk_start+1} to {chunk_end} of {total_paras}. Press Enter to continue... ---\n")
                    else:
                        raise EOFError  # force auto when not true TTY
                except EOFError:
                    # Stdin not available; fall back to timed pause
                    time.sleep(1)

def show_format_for_all(file_path, chunk_size=20):
    from docx import Document
    doc = Document(file_path)
    total_paras = len(doc.paragraphs)
    output_lines = []
    
    # Clean output file at start
    with open('output_chunk_bakti.txt', 'w', encoding='utf-8') as f:
        f.write("")
    
    print("\n=== Paragraphs and Runs ===")
    for chunk_start in range(0, total_paras, chunk_size):
        chunk_end = min(chunk_start + chunk_size, total_paras)
        chunk_lines = []
        for para_idx in range(chunk_start, chunk_end):
            para = doc.paragraphs[para_idx]
            chunk_lines.append(f"\nParagraph {para_idx+1}")
            chunk_lines.append(f"  Text: {para.text}")
            chunk_lines.append(f"  Style: {para.style.name if para.style else 'None'}")
            adv_styles = get_paragraph_advanced_styles(para)
            for k, v in adv_styles.items():
                chunk_lines.append(f"    {k}: {v}")
            for run_idx, run in enumerate(para.runs):
                font = run.font
                clean_text = run.text if run.text else ""
                font_name = font.name if font.name else 'Default'
                font_size = font.size.pt if font.size else 'Default'
                bold = 'Yes' if run.bold else 'No'
                italic = 'Yes' if run.italic else 'No'
                underline = 'Yes' if run.underline else 'No'
                color = font.color.rgb if font.color and font.color.rgb else 'Default'
                chunk_lines.append(f"    Run {run_idx+1}:")
                chunk_lines.append(f"      Text      : {clean_text}")
                chunk_lines.append(f"      Font Name : {font_name}")
                chunk_lines.append(f"      Font Size : {font_size}")
                chunk_lines.append(f"      Bold      : {bold}")
                chunk_lines.append(f"      Italic    : {italic}")
                chunk_lines.append(f"      Underline : {underline}")
                chunk_lines.append(f"      Color     : {color}")
                adv = get_run_advanced_styles(run)
                for k, v in adv.items():
                    chunk_lines.append(f"      {k}: {v}")
        print("\n".join(chunk_lines))
        with open('output_chunk_bakti.txt', 'a', encoding='utf-8') as f:
            f.write("\n".join(chunk_lines) + "\n\n")
        if chunk_end < total_paras:
            # Get auto_stream from global scope
            auto_stream = getattr(args, 'auto_stream', True)
            
            if auto_stream:
                # Auto-streaming mode: brief delay between chunks
                time.sleep(0.5)
            else:
                try:
                    # Try interactive pause
                    if sys.stdin.isatty():
                        input(f"\n--- Showing paragraphs {chunk_start+1} to {chunk_end} of {total_paras}. Press Enter to continue... ---\n")
                    else:
                        raise EOFError  # force auto when not true TTY
                except EOFError:
                    # Stdin not available; fall back to timed pause
                    time.sleep(1)

    print("\n=== Tables ===")
    for tbl_idx, table in enumerate(doc.tables):
        print(f"\nTable {tbl_idx+1}:")
        total_cells = sum(len(row.cells) for row in table.rows)
        cell_count = 0
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_count += 1
                text = cell.text.replace('\n', ' | ')
                print(f"  Cell ({row_idx+1},{cell_idx+1}): {text}")
                # Show formatting for each paragraph in cell
                for para in cell.paragraphs:
                    print(f"    Cell Paragraph: {para.text}")
                    print(f"      Style: {para.style.name if para.style else 'None'}")
                    adv_styles = get_paragraph_advanced_styles(para)
                    for k, v in adv_styles.items():
                        print(f"        {k}: {v}")
                    for run_idx, run in enumerate(para.runs):
                        font = run.font
                        font_name = font.name if font.name else 'Default'
                        font_size = font.size.pt if font.size else 'Default'
                        bold = 'Yes' if run.bold else 'No'
                        italic = 'Yes' if run.italic else 'No'
                        underline = 'Yes' if run.underline else 'No'
                        color = font.color.rgb if font.color and font.color.rgb else 'Default'
                        print(f"        Run {run_idx+1}:")
                        print(f"          Text      : {run.text}")
                        print(f"          Font Name : {font_name}")
                        print(f"          Font Size : {font_size}")
                        print(f"          Bold      : {bold}")
                        print(f"          Italic    : {italic}")
                        print(f"          Underline : {underline}")
                        print(f"          Color     : {color}")
                        adv = get_run_advanced_styles(run)
                        for k, v in adv.items():
                            print(f"          {k}: {v}")
                # Chunk tables by cell
                if chunk_size > 0 and cell_count % chunk_size == 0 and cell_count < total_cells:
                    input(f"\n--- Showing table cells {cell_count-chunk_size+1} to {cell_count} of {total_cells} in Table {tbl_idx+1}. Press Enter to continue... ---\n")

import argparse

def show_toc_fields(file_path):
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(file_path)
    print("Table of Contents Fields:")
    found = False
    for para in doc.paragraphs:
        p = para._p
        for fld in p.iter():
            if fld.tag.endswith('fldSimple') and 'TOC' in fld.attrib.get(qn('w:instr'), ''):
                print(f"  Paragraph: '{para.text}' (ToC field)")
                found = True
    if not found:
        print("  No ToC fields found.")

def show_headings(file_path):
    doc = Document(file_path)
    print("Headings:")
    found = False
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading'):
            print(f"  Paragraph {i+1}: '{para.text}' (Style: {para.style.name})")
            found = True
    if not found:
        print("  No headings found.")

def show_captions(file_path):
    doc = Document(file_path)
    print("Captions:")
    found = False
    for i, para in enumerate(doc.paragraphs):
        if para.style and 'caption' in para.style.name.lower():
            print(f"  Paragraph {i+1}: '{para.text}' (Style: {para.style.name})")
            found = True
        elif para.text.strip().lower().startswith(('figure', 'table')):
            print(f"  Paragraph {i+1}: '{para.text}' (Possible caption)")
            found = True
    if not found:
        print("  No captions found.")

def show_page_breaks(file_path):
    from docx import Document
    doc = Document(file_path)
    print("Page Breaks:")
    found = False
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            if 'pageBreak' in run._element.xml:
                print(f"  Paragraph {i+1}: '{para.text}' (Page break)")
                found = True
    if not found:
        print("  No page breaks found.")

def show_header_footer_and_layout(file_path):
    from docx import Document
    doc = Document(file_path)
    section = doc.sections[0]
    print("Header/Footer and Page Layout:")
    # Header
    print("Header:")
    for para in section.header.paragraphs:
        print(f"  {para.text}")
    # Footer
    print("Footer:")
    for para in section.footer.paragraphs:
        print(f"  {para.text}")
    # Layout
    print(f"Page size: {section.page_width} x {section.page_height}")
    print(f"Orientation: {'Landscape' if section.page_width > section.page_height else 'Portrait'}")
    print(f"Margins (cm): top={section.top_margin/360000.0:.2f}, right={section.right_margin/360000.0:.2f}, bottom={section.bottom_margin/360000.0:.2f}, left={section.left_margin/360000.0:.2f}")

def show_academic_structure(file_path):
    print("\n=== Academic Structure Summary ===")
    show_toc_fields(file_path)
    show_headings(file_path)
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Academic DOCX Reader")
    parser.add_argument('file_path', help='Path to DOCX file')
    parser.add_argument('--extract-images', action='store_true', help='Extract images from DOCX')
    parser.add_argument('--extract-content', action='store_true', help='Extract content and styles (chunked)')
    parser.add_argument('--chunk-size', type=int, default=20, help='Chunk size for content extraction')
    # Academic features
    parser.add_argument('--show-structure', action='store_true', help='Show academic structure summary')
    parser.add_argument('--show-toc', action='store_true', help='Show Table of Contents fields')
    parser.add_argument('--show-headings', action='store_true', help='Show all headings')
    parser.add_argument('--show-captions', action='store_true', help='Show all captions')
    parser.add_argument('--show-page-breaks', action='store_true', help='Show page breaks')
    parser.add_argument('--show-header-footer', action='store_true', help='Show header, footer, and layout')
    parser.add_argument('--show-format', action='store_true', help='Show formatting for all text and tables')
    parser.add_argument('--format-chunk-size', type=int, default=20, help='Chunk size for --show-format (default: 20)')
    parser.add_argument('--auto-stream', action='store_true', help='Stream all output without pauses', default=True)
    args = parser.parse_args()

    if args.extract_images:
        extract_images_from_docx(args.file_path)
    elif args.extract_content:
        extract_content_and_style(args.file_path, chunk_size=args.chunk_size)
    elif args.show_structure:
        show_academic_structure(args.file_path)
    elif args.show_toc:
        show_toc_fields(args.file_path)
    elif args.show_headings:
        show_headings(args.file_path)
    elif args.show_captions:
        show_captions(args.file_path)
    elif args.show_page_breaks:
        show_page_breaks(args.file_path)
    elif args.show_header_footer:
        show_header_footer_and_layout(args.file_path)
    elif args.show_format:
        show_format_for_all(args.file_path, chunk_size=args.format_chunk_size)
    else:
        read_docx(args.file_path)
