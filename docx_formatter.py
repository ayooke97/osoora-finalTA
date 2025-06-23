from docx import Document
import sys
import os
import argparse

def standardize_document_format(input_file, output_file):
    """
    Standardizes the formatting of a DOCX file to ensure consistent output
    in read_docx.py that matches output_chunk.txt format.
    """
    print(f"Standardizing format for: {input_file}")
    doc = Document(input_file)
    
    # Process paragraphs
    for para in doc.paragraphs:
        # Ensure consistent line spacing (1.5 as in output_chunk.txt)
        para.paragraph_format.line_spacing = 1.5
        
        # Process runs within paragraphs
        for run in para.runs:
            # Set default font size if not specified
            if not run.font.size:
                # Standard size in output_chunk.txt is 12pt
                run.font.size = 12 * 12700  # 12pt in docx internal units
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    # Apply same formatting to table paragraphs
                    para.paragraph_format.line_spacing = 1.5
                    
                    for run in para.runs:
                        if not run.font.size:
                            run.font.size = 12 * 12700
    
    # Save the standardized document
    doc.save(output_file)
    print(f"Successfully standardized document format: {output_file}")
    return True

def main():
    parser = argparse.ArgumentParser(description='Standardize DOCX formatting to match output_chunk.txt format')
    parser.add_argument('input', help='Input DOCX file path')
    parser.add_argument('output', help='Output DOCX file path')
    args = parser.parse_args()
    
    if standardize_document_format(args.input, args.output):
        print("\nFormat standardization complete!")
        print("Run read_docx.py on the new file to verify consistent formatting.")

if __name__ == "__main__":
    main()
